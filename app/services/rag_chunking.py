"""Document text extraction, cleaning, and chunking for RAG ingestion.

Text/markdown/code are decoded directly (stdlib). PDF uses ``pypdf`` and DOCX
uses ``python-docx`` — both pure-Python and lazily imported, so the service
still loads if they are absent (an unsupported file then fails its own job
cleanly rather than breaking the app).
"""
from __future__ import annotations

import io
import os
import re
import unicodedata

# Extensions decoded as plain UTF-8 text.
_TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".mdx", ".rst", ".csv", ".tsv", ".json",
    ".log", ".yaml", ".yml", ".xml", ".html", ".htm", ".ini", ".cfg", ".toml",
    ".py", ".js", ".ts", ".tsx", ".jsx", ".java", ".c", ".cpp", ".h", ".hpp",
    ".cs", ".go", ".rs", ".rb", ".php", ".sh", ".sql", ".dart", ".kt", ".swift",
}
_PDF_EXTS = {".pdf"}
_DOCX_EXTS = {".docx"}

SUPPORTED_EXTS = _TEXT_EXTS | _PDF_EXTS | _DOCX_EXTS


def is_supported(filename: str) -> bool:
    return _ext(filename) in SUPPORTED_EXTS


def _ext(filename: str) -> str:
    return os.path.splitext(filename or "")[1].lower()


def extract_text(filename: str, content: bytes) -> str:
    """Extract raw text from an uploaded file's bytes. Raises ValueError with a
    user-facing message for unsupported or unreadable files."""
    ext = _ext(filename)
    if ext in _PDF_EXTS:
        return _extract_pdf(content)
    if ext in _DOCX_EXTS:
        return _extract_docx(content)
    if ext in _TEXT_EXTS:
        return content.decode("utf-8", errors="replace")
    # Unknown extension: accept it only if it decodes as mostly-printable text.
    text = content.decode("utf-8", errors="replace")
    if _looks_binary(text):
        raise ValueError(
            f"Unsupported file type '{ext or filename}'. Supported: PDF, DOCX, "
            f"and text/markdown/code files."
        )
    return text


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover
        raise ValueError("PDF support unavailable (pypdf not installed)") from exc
    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:
        raise ValueError(f"Could not read PDF: {exc}") from exc
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    text = "\n\n".join(parts).strip()
    if not text:
        raise ValueError(
            "No extractable text in this PDF (it may be scanned images — OCR "
            "is not supported)."
        )
    return text


def _extract_pdf_pages(content: bytes) -> list[str]:
    """PDF text split per page (page boundaries preserved for page numbers)."""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return pages


def extract_pages(filename: str, content: bytes) -> "Optional[list[str]]":
    """Per-page text for PDFs (so chunks can carry a page number); ``None`` for
    formats without a page concept (text/DOCX)."""
    if _ext(filename) in _PDF_EXTS:
        try:
            return _extract_pdf_pages(content)
        except Exception:
            return None
    return None


def build_paged_text(pages: list[str]) -> "tuple[str, list[tuple[int, int, int]]]":
    """Clean each page and concatenate; return (combined_text, spans) where each
    span is (char_start, char_end, page_number) into combined_text. Offsets match
    ``\\n\\n``.join(cleaned_pages) exactly so a chunk's char_start maps to a page."""
    JOIN = "\n\n"
    parts: list[str] = []
    spans: list[tuple[int, int, int]] = []
    offset = 0
    for i, raw in enumerate(pages, 1):
        ct = clean_text(raw)
        if not ct:
            continue
        start = offset
        parts.append(ct)
        offset += len(ct)
        spans.append((start, offset, i))
        offset += len(JOIN)     # the joiner that precedes the next page
    return JOIN.join(parts), spans


def page_for_offset(offset: int, spans: "list[tuple[int, int, int]]") -> "Optional[int]":
    """Which page a char offset falls in (nearest page for offsets in a joiner)."""
    for start, end, page in spans:
        if start <= offset < end:
            return page
    if spans and offset >= spans[-1][1]:
        return spans[-1][2]
    return spans[0][2] if spans else None


def _extract_docx(content: bytes) -> str:
    try:
        import docx  # python-docx
    except Exception as exc:  # pragma: no cover
        raise ValueError("DOCX support unavailable (python-docx not installed)") from exc
    try:
        document = docx.Document(io.BytesIO(content))
    except Exception as exc:
        raise ValueError(f"Could not read DOCX: {exc}") from exc
    lines = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    text = "\n".join(lines).strip()
    if not text:
        raise ValueError("No extractable text in this DOCX.")
    return text


def _looks_binary(text: str) -> bool:
    if not text:
        return True
    sample = text[:2000]
    # Many replacement chars or NULs → the bytes weren't really text.
    bad = sample.count("�") + sample.count("\x00")
    return bad > max(10, len(sample) * 0.1)


def clean_text(text: str) -> str:
    """Normalize whitespace and strip control characters while preserving
    paragraph structure (blank lines) so chunking can break on boundaries."""
    if not text:
        return ""
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Drop control chars except newline and tab.
    text = "".join(ch for ch in text if ch == "\n" or ch == "\t" or unicodedata.category(ch)[0] != "C")
    # Collapse runs of spaces/tabs; trim each line; cap consecutive blank lines.
    lines = [re.sub(r"[ \t]+", " ", ln).strip() for ln in text.split("\n")]
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


_BREAKS = ["\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " "]


def chunk_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    """Split cleaned text into overlapping chunks no larger than ``chunk_size``
    characters, preferring to cut on paragraph/sentence/word boundaries."""
    text = (text or "").strip()
    n = len(text)
    if n == 0:
        return []
    if n <= chunk_size:
        return [text]

    overlap = max(0, min(overlap, chunk_size // 2))
    chunks: list[str] = []
    start = 0
    while start < n:
        end = min(start + chunk_size, n)
        if end < n:
            window = text[start:end]
            cut = _best_break(window, int(chunk_size * 0.5))
            if cut > 0:
                end = start + cut
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= n:
            break
        start = max(end - overlap, start + 1)
    return chunks


def _best_break(window: str, min_index: int) -> int:
    """Index just past the best boundary in ``window`` at or after ``min_index``,
    or -1 if none — so we never cut absurdly early."""
    for pat in _BREAKS:
        idx = window.rfind(pat)
        if idx >= min_index:
            return idx + len(pat)
    return -1


# ─── Structure-aware + parent/child chunking ────────────────────────────────
# (docs/semantic-embedding/03-chunking.md). Detect heading structure, split each
# section into a large *parent* (returned to the LLM) and small *child* chunks
# (embedded + searched). Headings are prepended to chunk text — this measurably
# improves retrieval because the section context is embedded with the content.

_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*#*$")
_LABEL_HEADING_RE = re.compile(
    r"^(?:Title|Section|Subsection|Subheading|Chapter|Part|Step)\s*:\s*(.+)$", re.I
)


def _detect_heading(line: str) -> Optional[tuple[int, str]]:
    """Return (level, title) if the line is a heading, else None."""
    m = _MD_HEADING_RE.match(line)
    if m:
        return len(m.group(1)), m.group(2).strip()
    m = _LABEL_HEADING_RE.match(line)
    if m:
        return 3, m.group(1).strip()
    return None


def _split_sections(text: str, structure_aware: bool) -> list[tuple[list[str], str, int]]:
    """Split into (heading_path, body, char_start). A heading path is the stack
    of enclosing headings, e.g. ["Authentication", "Refresh Token Rotation"]."""
    if not structure_aware:
        return [([], text, 0)]
    sections: list[tuple[list[str], str, int]] = []
    stack: list[tuple[int, str]] = []
    buf: list[str] = []
    buf_start = 0
    pos = 0

    def flush() -> None:
        body = "\n".join(buf).strip()
        if body:
            sections.append(([t for _, t in stack], body, buf_start))

    for line in text.split("\n"):
        h = _detect_heading(line)
        if h is not None:
            flush()
            level, title = h
            while stack and stack[-1][0] >= level:
                stack.pop()
            stack.append((level, title))
            buf = []
            buf_start = pos + len(line) + 1
        else:
            if not buf:
                buf_start = pos
            buf.append(line)
        pos += len(line) + 1
    flush()
    return sections or [([], text, 0)]


def _windows(text: str, size: int, overlap: int) -> list[tuple[str, int]]:
    """Boundary-preferring sliding windows over ``text`` → [(chunk, char_offset)]."""
    text = text.strip()
    n = len(text)
    if n == 0:
        return []
    if n <= size:
        return [(text, 0)]
    overlap = max(0, min(overlap, size // 2))
    out: list[tuple[str, int]] = []
    start = 0
    while start < n:
        end = min(start + size, n)
        if end < n:
            cut = _best_break(text[start:end], int(size * 0.5))
            if cut > 0:
                end = start + cut
        piece = text[start:end].strip()
        if piece:
            out.append((piece, start))
        if end >= n:
            break
        start = max(end - overlap, start + 1)
    return out


def _with_heading(label: Optional[str], body: str) -> str:
    return f"{label}\n{body}" if label else body


def chunk_document(
    text: str,
    *,
    child_size: int,
    parent_size: int,
    overlap: int,
    structure_aware: bool = True,
) -> list[dict]:
    """Structure-aware parent/child chunking. Returns a flat list of CHILD
    chunks (the searchable units); each carries the enclosing parent text so
    ingestion can store parent rows and link children to them.

    Each item: {text, raw_text, section, char_start, char_end, parent_key,
                parent_text, parent_char_start, parent_char_end}
    ``parent_key`` groups children that share one parent. When parent_size <=
    child_size the parent == the child (flat chunking, no expansion)."""
    text = (text or "").strip()
    if not text:
        return []
    parent_size = max(parent_size, child_size)  # parent must be ≥ child
    out: list[dict] = []
    parent_key = 0
    for path, body, sec_start in _split_sections(text, structure_aware):
        label = " ▸ ".join(path) if path else None
        parents = _windows(body, parent_size, 0)
        for ptext, poff in parents:
            p_start = sec_start + poff
            p_full = _with_heading(label, ptext)
            children = _windows(ptext, child_size, overlap)
            for ctext, coff in children:
                c_start = p_start + coff
                out.append({
                    "text": _with_heading(label, ctext),   # embedded + stored
                    "raw_text": ctext,
                    "section": label,
                    "char_start": c_start,
                    "char_end": c_start + len(ctext),
                    "parent_key": parent_key,
                    "parent_text": p_full,
                    "parent_char_start": p_start,
                    "parent_char_end": p_start + len(ptext),
                })
            parent_key += 1
    return out
